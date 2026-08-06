import pytest

from search.job_store import save_jobs
from tailoring.applications import upsert_application
from tailoring.dossier import check_for_edits, sync_workspace_documents


@pytest.fixture
def job(isolated_data):
    import profile.storage as storage
    storage.save_profile({"name": "Jane Doe"})

    j = {
        "source": "linkedin", "job_id": "1", "title": "Head of IT",
        "organization": "Acme Corp", "location": "Remote",
    }
    save_jobs([j])
    return j


def test_unedited_cover_letter_is_not_reported_as_changed(job):
    # Regression for a real bug (2026-08-04, Zahir hit this live): a cover
    # letter he never touched was reported as "4 lines added, 6 removed".
    # Root cause - check_for_edits() diffed the RAW drafted text against
    # text extracted back out of the rendered .docx, but rendering a cover
    # letter always injects a date/company/address block that never existed
    # in the raw text, and strips blank lines - guaranteeing a spurious
    # diff on every single cover letter, edited or not.
    body = "Dear Hiring Team,\n\nI am excited to apply for this role.\n\nSincerely,\nJane Doe"
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)

    result = check_for_edits("linkedin", "1")
    assert "no_workspace_file" not in result["cover_letter"]
    assert result["cover_letter"]["changed"] is False
    assert result["cover_letter"]["diff"] == []


def test_unedited_resume_with_bullets_is_not_reported_as_changed(job):
    body = (
        "Jane Doe\n"
        "jane@example.com\n"
        "\n"
        "PROFESSIONAL EXPERIENCE\n"
        "- Did a thing.\n"
        "- Did another thing.\n"
    )
    upsert_application("linkedin", "1", status="under review", resume_text=body)
    sync_workspace_documents("linkedin", "1", ["resume"], {"resume": body}, {"name": "Jane Doe"}, job)

    result = check_for_edits("linkedin", "1")
    assert "no_workspace_file" not in result["resume"]
    assert result["resume"]["changed"] is False
    assert result["resume"]["diff"] == []


def test_actually_edited_cover_letter_is_still_detected(job):
    body = "Dear Hiring Team,\n\nI am excited to apply.\n\nSincerely,\nJane Doe"
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)

    from tailoring.dossier import dossier_dir, _workspace_filename
    from docx import Document

    folder = dossier_dir("linkedin", "1", "Acme Corp", "Head of IT")
    file_path = folder / _workspace_filename("cover_letter_text", "Jane Doe", job)
    doc = Document(str(file_path))
    doc.paragraphs[-1].add_run(" - added by hand in Word.")
    doc.save(str(file_path))

    result = check_for_edits("linkedin", "1")
    assert result["cover_letter"]["changed"] is True
    assert result["cover_letter"]["diff"]


def test_sync_migrates_old_descriptive_filename_without_losing_edits(job):
    # Integration-level version of the unit tests in test_dossier_naming.py:
    # a folder created under the pre-2026-08-05 {Name}_{DocType}_{Title}_
    # {Organization}.docx format (before the 50-char ceiling existed) must
    # have its real file found and renamed forward by the actual
    # sync_workspace_documents() call path a real regenerate goes through -
    # not just by the migration helper in isolation - so Zahir's genuinely
    # edited copy is never orphaned under the old filename.
    from docx import Document

    from tailoring.dossier import _legacy_descriptive_workspace_filename, _workspace_filename, dossier_dir

    body = "Dear Hiring Team,\n\nI am excited to apply.\n\nSincerely,\nJane Doe"
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)

    # Simulate a folder created under the old format: draft once (creates a
    # real .docx), rename it onto the old descriptive filename, then edit it
    # in place - exactly like test_actually_edited_cover_letter_is_still_
    # detected does. A fake plain-text file wouldn't open as a real .docx,
    # so extraction would fail and skip the backup step entirely, silently
    # defeating the very thing this test is meant to verify.
    folder = dossier_dir("linkedin", "1", "Acme Corp", "Head of IT")
    old_name = _legacy_descriptive_workspace_filename("cover_letter_text", "Jane Doe", job)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)
    new_name = _workspace_filename("cover_letter_text", "Jane Doe", job)
    (folder / new_name).rename(folder / old_name)

    doc = Document(str(folder / old_name))
    doc.paragraphs[-1].add_run(" - added by hand in Word.")
    doc.save(str(folder / old_name))

    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)

    assert not (folder / old_name).exists()
    assert (folder / new_name).exists()
    # The migrated file's content differed from the freshly-rendered
    # baseline (it held Jane's real edit, not a matching draft) - sync's
    # own edit-detection must treat the just-migrated file the same as any
    # other on-disk file and back up the edit rather than silently
    # overwriting it with the fresh regenerate.
    backups = [p for p in folder.iterdir() if ".edited-" in p.name]
    assert len(backups) == 1
    backup_text = "\n".join(p.text for p in Document(str(backups[0])).paragraphs)
    assert "added by hand in Word" in backup_text


def test_sync_migrates_old_hash_suffixed_filename_without_losing_edits(job):
    # Same integration-level guarantee as the descriptive-format migration
    # test above, one format later: a folder created under the brief
    # 2026-08-05 {Name}_{DocType}_{Organization}_{hash}.docx format (commit
    # 51120a7, superseded almost immediately when Zahir asked for the
    # trailing hex string removed) must have its real file found and
    # renamed forward by the real sync_workspace_documents() call path, not
    # just by the migration helper in isolation.
    from docx import Document

    from tailoring.dossier import _legacy_hash_suffixed_workspace_filename, _workspace_filename, dossier_dir

    body = "Dear Hiring Team,\n\nI am excited to apply.\n\nSincerely,\nJane Doe"
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)

    folder = dossier_dir("linkedin", "1", "Acme Corp", "Head of IT")
    old_name = _legacy_hash_suffixed_workspace_filename("cover_letter_text", "Jane Doe", job)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)
    new_name = _workspace_filename("cover_letter_text", "Jane Doe", job)
    (folder / new_name).rename(folder / old_name)

    doc = Document(str(folder / old_name))
    doc.paragraphs[-1].add_run(" - added by hand in Word.")
    doc.save(str(folder / old_name))

    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)

    assert not (folder / old_name).exists()
    assert (folder / new_name).exists()
    backups = [p for p in folder.iterdir() if ".edited-" in p.name]
    assert len(backups) == 1
    backup_text = "\n".join(p.text for p in Document(str(backups[0])).paragraphs)
    assert "added by hand in Word" in backup_text


def test_regenerate_of_untouched_document_does_not_spam_edited_backup(job):
    # sync_workspace_documents() had the same raw-vs-rendered comparison bug
    # in its own edit-detection - every regenerate of an untouched document
    # was silently creating an unnecessary .edited-<timestamp>.docx backup.
    body = "Dear Hiring Team,\n\nI am excited to apply.\n\nSincerely,\nJane Doe"
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)
    upsert_application("linkedin", "1", status="under review", cover_letter_text=body)
    sync_workspace_documents("linkedin", "1", ["cover_letter"], {"cover_letter": body}, {"name": "Jane Doe"}, job)

    from tailoring.dossier import dossier_dir
    folder = dossier_dir("linkedin", "1", "Acme Corp", "Head of IT")
    backups = [p for p in folder.iterdir() if ".edited-" in p.name]
    assert backups == []
