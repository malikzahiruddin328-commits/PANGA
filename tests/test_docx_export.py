import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from tailoring.docx_export import NAME_ACCENT_COLOR, NAME_SIZE, BODY_FONT, BODY_SIZE, text_to_docx_bytes

SAMPLE_RESUME = (
    "Jane Doe\n"
    "jane@example.com | (555) 123-4567 | linkedin.com/in/janedoe\n"
    "\n"
    "PROFESSIONAL SUMMARY\n"
    "Some summary text here.\n"
    "\n"
    "PROFESSIONAL EXPERIENCE\n"
    "Acme Corp  01/2020 - Present\n"
    "- Did a thing.\n"
)


def _load(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_name_is_center_aligned_like_the_base_template():
    # Regression for a real bug (2026-08-04): the name paragraph was styled
    # bold/large/colored but never actually centered, even though the
    # contact-info line right under it always was - so a generated resume's
    # name sat flush-left while everything else matched Zahir's real resume
    # template it was ported from.
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    name_para = doc.paragraphs[0]
    assert name_para.text == "Jane Doe"
    assert name_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_name_run_is_bold_large_and_accent_colored():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    run = doc.paragraphs[0].runs[0]
    assert run.bold is True
    assert run.font.size == NAME_SIZE
    assert run.font.color.rgb == NAME_ACCENT_COLOR


def test_contact_line_is_center_aligned():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    contact_para = doc.paragraphs[1]
    assert "jane@example.com" in contact_para.text
    assert contact_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_name_size_scales_with_body_size_pt_and_stays_centered():
    # USAJOBS resumes shrink body text to fit the 2-page hard cap
    # (tailoring/dossier.py's body_size_pt override) - the name should
    # shrink proportionally and stay centered, not revert to a default.
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME, body_size_pt=9.0))
    name_para = doc.paragraphs[0]
    run = name_para.runs[0]
    assert run.font.size.pt == round(9.0 * (NAME_SIZE.pt / 10.5))
    assert name_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_base_font_and_size_applied_to_normal_style():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    normal = doc.styles["Normal"]
    assert normal.font.name == BODY_FONT
    assert normal.font.size == BODY_SIZE


def test_section_header_is_bold_and_accent_colored_not_centered():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    header_para = next(p for p in doc.paragraphs if p.text == "PROFESSIONAL SUMMARY")
    run = header_para.runs[0]
    assert run.bold is True
    assert run.font.color.rgb == NAME_ACCENT_COLOR
    assert header_para.alignment != WD_ALIGN_PARAGRAPH.CENTER


def test_section_header_carries_a_real_heading_style():
    # Real ask 2026-08-06 (Mirror/Zahir): some ATS/recruiter tools detect
    # section boundaries from paragraph style metadata, not visual weight
    # alone - bold body text wasn't enough, needs a real Word heading
    # style underneath the existing bold/color/compact-size look.
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    header_para = next(p for p in doc.paragraphs if p.text == "PROFESSIONAL SUMMARY")
    assert header_para.style.name == "Heading 2"
    run = header_para.runs[0]
    assert run.font.size == BODY_SIZE
    assert run.font.name == BODY_FONT


def test_date_range_line_is_bold_not_centered():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    role_para = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert role_para.runs[0].bold is True
    assert role_para.alignment != WD_ALIGN_PARAGRAPH.CENTER


def test_date_range_in_the_real_month_yyyy_format_still_matches_and_right_aligns():
    # Real bug found 2026-08-06 (Mirror/Zahir): the old date-range regex
    # required only whitespace between the dash and the closing year, so
    # it never matched RESUME_SPEC's own requested "Month YYYY - Month
    # YYYY" format (the month name in the middle broke it) - every date
    # line in a real generated resume fell through to a plain, unbolded,
    # unaligned paragraph instead of this branch. This is the format the
    # app actually produces, not the MM/YYYY shorthand in SAMPLE_RESUME.
    text = (
        "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\n"
        "Acme Corp - Springfield, IL \t\t September 2018 - January 2026\n"
        "- Did a thing.\n"
    )
    doc = _load(text_to_docx_bytes(text))
    role_para = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert [r.text for r in role_para.runs] == [
        "Acme Corp - Springfield, IL",
        "\t",
        "September 2018 - January 2026",
    ]
    assert all(r.bold for r in role_para.runs)
    tab_stops = list(role_para.paragraph_format.tab_stops)
    assert len(tab_stops) == 1
    assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT


def test_date_range_ending_in_present_still_right_aligns():
    text = "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\nAcme Corp June 2020 - Present\n"
    doc = _load(text_to_docx_bytes(text))
    role_para = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert role_para.runs[-1].text == "June 2020 - Present"
    assert len(list(role_para.paragraph_format.tab_stops)) == 1


def test_en_dash_date_separator_normalized_to_plain_hyphen():
    # Real ask 2026-08-06 (Mirror/Zahir): a few older ATS parsers mis-
    # tokenize en-dashes - the rendered separator must always be a plain
    # hyphen regardless of what the drafted text used.
    text = "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\nAcme Corp September 2018 – January 2026\n"
    doc = _load(text_to_docx_bytes(text))
    role_para = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert role_para.runs[-1].text == "September 2018 - January 2026"


def test_to_date_separator_normalized_to_plain_hyphen():
    text = "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\nAcme Corp September 2018 to January 2026\n"
    doc = _load(text_to_docx_bytes(text))
    role_para = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert role_para.runs[-1].text == "September 2018 - January 2026"


def test_date_range_on_its_own_line_merges_onto_the_preceding_title_line():
    # Real bug found live 2026-08-06, verifying the VP-tier title fix
    # against an actual generation: some drafts put the date range on its
    # OWN line, right after the title/company line, instead of sharing it
    # ("Head of IT\nJanuary 2024 - January 2026\n"). Before this fix, the
    # empty prefix rendered as its own near-blank paragraph - a lone
    # unbolded "Head of IT" line, then a date floating flush-right on the
    # next line with nothing visibly tying it to that title.
    text = "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\nHead of IT\nJanuary 2024 - January 2026\n- Did a thing.\n"
    doc = _load(text_to_docx_bytes(text))
    role_para = next(p for p in doc.paragraphs if "Head of IT" in p.text)
    assert [r.text for r in role_para.runs] == ["Head of IT", "\t", "January 2024 - January 2026"]
    assert all(r.bold for r in role_para.runs)
    assert len(list(role_para.paragraph_format.tab_stops)) == 1
    # Only one paragraph for this role's title+date, not two.
    assert sum(1 for p in doc.paragraphs if "Head of IT" in p.text or "January 2024" in p.text) == 1


def test_date_range_on_its_own_line_after_a_bullet_does_not_merge_into_it():
    # A lone date-range line straight after a bullet (not a title/company
    # line) shouldn't merge into that bullet's "List Bullet" paragraph -
    # only a plain "Normal"-style preceding paragraph is a safe merge
    # target.
    text = "Jane Doe\njane@example.com\n\nPROFESSIONAL EXPERIENCE\n- Did a thing.\nJanuary 2024 - January 2026\n"
    doc = _load(text_to_docx_bytes(text))
    bullet_para = next(p for p in doc.paragraphs if p.text == "Did a thing.")
    assert bullet_para.style.name == "List Bullet"
    date_para = next(p for p in doc.paragraphs if "January 2024" in p.text)
    assert date_para is not bullet_para


def test_all_caps_name_is_rendered_in_title_case():
    # Real bug found 2026-08-06 (Mirror/Zahir): a drafted resume sometimes
    # echoes the candidate's name in whatever casing the source resume
    # used, including ALL CAPS - a documented ATS-parsing problem, not
    # just style. Must be normalized regardless of the drafted text's
    # casing, not left to prompt compliance alone.
    text = "JANE DOE\njane@example.com\n"
    doc = _load(text_to_docx_bytes(text))
    name_para = doc.paragraphs[0]
    assert name_para.runs[0].text == "Jane Doe"


def test_mixed_case_name_is_left_untouched():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    assert doc.paragraphs[0].runs[0].text == "Jane Doe"


def test_bullet_line_uses_list_bullet_style():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    bullet_para = next(p for p in doc.paragraphs if p.text == "Did a thing.")
    assert bullet_para.style.name == "List Bullet"


def test_author_sets_core_properties_instead_of_python_docx_default():
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME, author="Jane Doe"))
    assert doc.core_properties.author == "Jane Doe"
    assert doc.core_properties.last_modified_by == "Jane Doe"


def test_contact_info_is_never_rendered_as_a_real_hyperlink():
    # Confirmed correct 2026-08-06 (Mirror/Zahir's ATS-parser review):
    # plain text is the safer choice for ATS extraction of email/phone/
    # LinkedIn - this asserts it stays that way rather than someone later
    # "improving" it into real w:hyperlink field objects.
    doc = _load(text_to_docx_bytes(SAMPLE_RESUME))
    assert "<w:hyperlink" not in doc.element.xml
