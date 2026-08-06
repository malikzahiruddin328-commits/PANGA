"""Converts a drafted document's plain text into a real .docx file Zahir
can attach to an actual application - most application portals require an
uploaded file, not pasted text (Zahir's explicit request 2026-07-30: "these
need to be attached"). Generated fresh in memory each time a download is
requested, never written to disk, so the underlying drafted text can stay
encrypted at rest in applications.json like everything else in data/
without this needing its own plaintext-on-disk exception.

Styled to match Zahir's own resume conventions (2026-07-31 request: "the
document save function should follow my resume style") - inspected directly
from his real resume file (C:/Users/User/Desktop/GAND/Zahir Resume.docx) via
python-docx rather than guessed: Times New Roman throughout, a compact 10pt
body (headers are distinguished by bold, not by being bigger - an
information-dense style, not a spacious one), and his own teal accent color
on his name. Plain-text structure is inferred with simple heuristics (first
line = name, second = contact info, ALL-CAPS lines = section headers, "- "
lines = bullets, lines matching a date range = bold company/role lines) -
this matches what tailoring/drafting.py's RESUME_SPEC already instructs
Claude to produce, so the heuristics aren't guessing at arbitrary text, they
're reading a format this app itself asked for.
"""

import io
import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor

NAME_ACCENT_COLOR = RGBColor(0x00, 0x78, 0x6C)  # Zahir's own teal, from his real resume's Title run
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(20)

_MONTH_RE = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?\s+"
)
_DATE_TOKEN_RE = rf"(?:{_MONTH_RE})?(?:19|20)\d{{2}}"
# Anchored to the END of the line - a date range is always the trailing
# part of a company/title line, never something that can appear mid-
# sentence. Real gap found 2026-08-06 (Mirror/Zahir): the old version of
# this regex required only whitespace between the dash and the closing
# year, so it silently never matched RESUME_SPEC's own requested "Month
# YYYY - Month YYYY" format (the month name in the middle failed \s*) -
# every date-range line in a real generated resume fell through to a
# plain, unbolded, non-right-aligned paragraph instead of this branch.
_DATE_RANGE_RE = re.compile(
    rf"(?P<start>{_DATE_TOKEN_RE})\s*(?:-|–|to)\s*(?P<end>{_DATE_TOKEN_RE}|present)\s*$",
    re.IGNORECASE,
)


def _right_tab_position(doc: Document):
    """Distance from the left margin to the right margin - where a
    right-aligned tab stop needs to sit so text after it lands flush
    against the right edge of the page regardless of company-name length."""
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _normalize_name_case(line: str) -> str:
    """Real complaint 2026-08-06 (Zahir, via Mirror): a generated resume's
    name rendered in ALL CAPS - a documented ATS-compatibility problem
    (auto-parsers can mis-split all-caps names), not just a style
    preference. The drafted text sometimes echoes the name in whatever
    casing the candidate's original source resume used (RESUME_SPEC's
    "never invent/embellish" instruction means the AI won't normalize this
    on its own), so this can't be left to prompt compliance alone - force
    title case here regardless of what casing the drafted text contains."""
    if line == line.upper() and any(c.isalpha() for c in line):
        return line.title()
    return line


def _looks_like_header(line: str) -> bool:
    """ALL-CAPS section headers (PROFESSIONAL SUMMARY, EDUCATION, ...) - the
    literal convention tailoring/drafting.py's RESUME_SPEC instructs."""
    return bool(line) and line == line.upper() and any(c.isalpha() for c in line) and len(line) < 60


def _set_base_font(doc: Document, body_size: "Pt" = BODY_SIZE) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = body_size
    # A blank line in the source text used to ALSO get its own empty
    # paragraph on top of this - the two stacked, making the document look
    # sparse ("too many spaces/carriage returns", Zahir's real complaint on
    # the actual downloaded file). Space-after alone gives clean separation
    # between paragraphs without a manual blank paragraph in between.
    normal.paragraph_format.space_after = Pt(4)


def text_to_docx_bytes(text: str, author: str | None = None, body_size_pt: float | None = None) -> bytes:
    """body_size_pt overrides the default 10.5pt body size - used to shrink
    USAJOBS resumes (10pt, or 9pt if still running long) to fit its hard
    2-page upload limit (Zahir's explicit ask 2026-08-03: "reduce the size
    of text to 10 or 9 ... headings needs to be adjusted accordingly").
    Name/header sizing scales down proportionally with it so the document
    keeps the same visual hierarchy at any size, not just a smaller body
    next to an unchanged oversized name."""
    body_size = Pt(body_size_pt) if body_size_pt else BODY_SIZE
    name_size = Pt(round((body_size_pt or 10.5) * (NAME_SIZE.pt / 10.5)))

    doc = Document()
    _set_base_font(doc, body_size)

    # python-docx's default Word file properties (Save As dialog, File >
    # Info) list the author as literally "python-docx" - a real, visible
    # tell that the file was generated by a script, and the opposite of an
    # ATS-perfect, looks-hand-made document. Zahir's explicit request
    # 2026-07-31: it should say his own name there instead.
    if author:
        doc.core_properties.author = author
        doc.core_properties.last_modified_by = author

    lines = text.split("\n")
    seen_name = False
    in_contact_block = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            # Skip blank lines entirely rather than emitting an empty
            # paragraph for each one - Normal's space_after (above) already
            # gives every real paragraph breathing room, so an ALSO-blank
            # paragraph on top of that doubled the visual gap.
            in_contact_block = False
            continue

        if not seen_name:
            # First non-blank line: the candidate's name, styled like the
            # Title run in Zahir's own resume - bold, larger, his own accent
            # color rather than a generic default.
            p = doc.add_paragraph()
            run = p.add_run(_normalize_name_case(line))
            run.bold = True
            run.font.size = name_size
            run.font.color.rgb = NAME_ACCENT_COLOR
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            seen_name = True
            in_contact_block = True
            continue

        if in_contact_block and not _looks_like_header(line) and not line.startswith("- "):
            # Contact block right under the name - can be one combined line
            # or several (phone/email/LinkedIn/work authorization each on
            # their own line, seen in real generations), centered either
            # way like the contact line in his resume. Ends at the first
            # blank line, header, or bullet. Deliberately plain text, not a
            # w:hyperlink field (confirmed correct 2026-08-06, Mirror/
            # Zahir's ATS-parser review) - plain text is the safer choice
            # for ATS extraction of email/phone/LinkedIn; don't "improve"
            # this into real hyperlink objects later.
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        in_contact_block = False

        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue

        if _looks_like_header(line):
            # Bold, not bigger - his resume distinguishes headers by weight,
            # not size, keeping the document compact and information-dense.
            # Colored to match his name (his explicit request 2026-07-31),
            # a deliberate departure from his own resume file (whose actual
            # headers are plain black) - not a mismatch, an update he asked
            # for on top of it. Also tagged with the real "Heading 2" style,
            # not just visual bold (Mirror/Zahir, 2026-08-06): some ATS/
            # recruiter tools detect section boundaries from paragraph
            # style metadata, not visual weight alone - the run-level
            # overrides below keep the existing compact look on top of it.
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(line)
            run.bold = True
            run.font.name = BODY_FONT
            run.font.size = body_size
            run.font.color.rgb = NAME_ACCENT_COLOR
            continue

        date_match = _DATE_RANGE_RE.search(line) if len(line) < 120 else None
        if date_match:
            # Company/role lines carrying a date range - bold, and the date
            # itself right-aligned against the page margin regardless of
            # how long the company/title text before it is, matching the
            # "SK Life Science, Inc. (SKLSI)      09/2018 - 01/2026" layout
            # in his real resume (a real tab stop, not whitespace the AI
            # tried to pad with - literal spaces/tabs in the drafted text
            # don't reliably line up since this is a proportional font).
            prefix = line[: date_match.start()].rstrip(" \t")
            # Separator normalized to a plain hyphen regardless of whether
            # the drafted text used an en-dash or "to" (Mirror/Zahir,
            # 2026-08-06): a few older ATS parsers mis-tokenize en-dashes.
            date_part = f"{date_match.group('start')} - {date_match.group('end')}"
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(_right_tab_position(doc), WD_TAB_ALIGNMENT.RIGHT)
            p.add_run(prefix).bold = True
            p.add_run("\t").bold = True
            p.add_run(date_part).bold = True
            continue

        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_UNUSABLE_COMPANY_NAMES = {"", "confidential", "unknown", "n/a", "none"}


def cover_letter_to_docx_bytes(
    body_text: str,
    company_name: str | None = None,
    company_address: str | None = None,
    author: str | None = None,
) -> bytes:
    """Cover letters get their own renderer rather than reusing
    text_to_docx_bytes() above - that one always treats a document's first
    non-blank line as the candidate's name and renders it at 20pt, which is
    exactly right for a resume but meant a cover letter's opening line
    ("Dear Hiring Team,") was rendering as a giant 20pt heading (Zahir's
    real complaint on the actual downloaded file, 2026-08-01). A cover
    letter is a standard business letter, not a resume: today's date and a
    recipient block come first, then plain paragraphs at one consistent
    size throughout - no line in a cover letter should ever be styled as
    a "name".

    company_name/company_address come from the job posting, which usually
    only has a company name (job_store.py's schema has no address field at
    all) - genuinely unknown values fall back to a bracketed placeholder
    rather than a guess, since this becomes a real submitted document.
    """
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(8)

    if author:
        doc.core_properties.author = author
        doc.core_properties.last_modified_by = author

    today = date.today()
    doc.add_paragraph(f"{today:%B} {today.day}, {today:%Y}")

    clean_company = (company_name or "").strip()
    company_p = doc.add_paragraph()
    company_p.add_run(
        company_name if clean_company.lower() not in _UNUSABLE_COMPANY_NAMES else "[Company Name]"
    ).bold = True

    clean_address = (company_address or "").strip()
    address_p = doc.add_paragraph(clean_address if clean_address else "[Company Address]")
    address_p.paragraph_format.space_after = Pt(20)

    for raw_line in body_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        if author and line == author:
            # The closing signature line - bold, matching how the name
            # reads as a distinct element without blowing it up to resume
            # size the way the old shared renderer's heuristic did.
            run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
