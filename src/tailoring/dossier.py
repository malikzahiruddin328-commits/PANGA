"""Per-application workspace (Zahir's request 2026-07-30, extended
2026-07-31 into a real folder per application): one traceable folder per
job actually engaged with (applied, rejected, under review, etc.),
consolidating the posting itself, application status/timeline, drafted
documents, and interview prep into a single place - instead of having to
cross-reference jobs.json/applications.json/interview_prep.json by hand for
context on one job.

Two kinds of content live in this folder, written by two different
functions:
- dossier.md (write_dossier(), below) - a readable summary, regenerated
  wholesale from the source-of-truth stores every time any of them changes
  for a given job (see call sites in applications.py/interview_prep.py).
  Read-only output - nothing reads it back in, and nothing Zahir is meant
  to hand-edit.
- The actual editable documents - resume.docx, cover_letter.docx, etc.
  (sync_workspace_documents(), below) - written ONLY at the moment a
  document is freshly (re)generated, never as a side effect of an unrelated
  dossier refresh (status change, strategy tag, skip reason...), because
  Zahir is expected to open and edit these directly in Word. Rewriting them
  on every unrelated save would silently clobber his in-progress edits.

Plain files on disk, not encrypted (Zahir's explicit call 2026-07-31 - this
folder needs to be directly double-clickable/editable in Word, which
encryption at rest would prevent). Lives under data/ like every other Panga
store, so it persists across code updates/reinstalls the same way
jobs.json/applications.json do - the only way to lose it is deleting the
data/ folder itself, same as everything else here.
"""

import difflib
import hashlib
import io
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from profile.storage import load_profile
from search.job_store import load_jobs
from tailoring.applications import get_application
from tailoring.interview_prep import get_interview_prep
from tailoring.docx_export import text_to_docx_bytes, cover_letter_to_docx_bytes
from tailoring.unconfirmed_claims import has_unconfirmed_marker

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOSSIER_DIR = PROJECT_ROOT / "data" / "applications" / "dossiers"

DOC_LABELS = {
    "resume_text": "Resume",
    "cover_letter_text": "Cover letter",
    "exec_bio_text": "Executive bio",
    "leadership_summary_text": "Leadership summary",
}

# The generic names workspace files used before 2026-07-31 - kept only so
# _migrate_legacy_filename() can find and rename a folder's existing files
# the first time it's touched after the naming convention below changed.
_LEGACY_WORKSPACE_FILENAMES = {
    "resume_text": "resume.docx",
    "cover_letter_text": "cover_letter.docx",
    "exec_bio_text": "exec_bio.docx",
    "leadership_summary_text": "leadership_summary.docx",
}

# Hard ceiling Zahir needs (2026-08-05): some destinations he pastes/copies
# these files into truncate or mangle names past this length - a real
# posting title alone ("Senior Director of IT Enterprise Applications") can
# already run 46 characters, so Title was dropped from the filename
# entirely below (still preserved in the parent dossier folder name via
# _slug() - nothing is lost, just not duplicated).
_MAX_WORKSPACE_FILENAME_LEN = 50
_WORKSPACE_FILENAME_SUFFIX_LEN = 6

DOC_KEY_TO_FIELD = {
    "resume": "resume_text",
    "cover_letter": "cover_letter_text",
    "exec_bio": "exec_bio_text",
    "leadership_summary": "leadership_summary_text",
}


def _job_hash(source: str, job_id: str, length: int = 8) -> str:
    """Stable per-job identifier, independent of what job_id itself looks
    like (some sources store a full posting URL as job_id - see
    search/industry_boards.py) - used both for the dossier folder's own
    slug below and as the workspace filename's uniqueness suffix, so two
    different jobs never collide even when Name+DocType+Organization are
    otherwise identical (e.g. two different roles at the same company)."""
    return hashlib.sha1(f"{source}:{job_id}".encode()).hexdigest()[:length]


def _slug(source: str, job_id: str, organization: str | None, title: str | None) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", f"{organization or ''}-{title or ''}".lower()).strip("-")[:60]
    return f"{base or 'job'}-{_job_hash(source, job_id)}"


def dossier_dir(source: str, job_id: str, organization: str | None = None, title: str | None = None) -> Path:
    return DOSSIER_DIR / _slug(source, job_id, organization, title)


def dossier_path(source: str, job_id: str, organization: str | None = None, title: str | None = None) -> Path:
    return dossier_dir(source, job_id, organization, title) / "dossier.md"


def _sanitize_filename_part(value: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', "", value).strip().replace(" ", "_")


def _legacy_descriptive_workspace_filename(field: str, candidate_name: str | None, job: dict) -> str:
    """The {Name}_{DocType}_{Title}_{Organization}.docx format used
    2026-07-31 through 2026-08-05, before the 50-char ceiling below existed
    - kept only so _migrate_descriptive_filename() can find and rename a
    folder's existing files under this now-superseded format. A real
    posting title alone ("Senior Director of IT Enterprise Applications")
    could already run this past 100+ characters, which is what the new
    format below fixes."""
    doc_label = DOC_LABELS.get(field, field)
    parts = [candidate_name, doc_label, job.get("title"), job.get("organization")]
    safe_parts = [_sanitize_filename_part(p) for p in parts if p]
    return "_".join(safe_parts)[:150] + ".docx"


def _legacy_hash_suffixed_workspace_filename(field: str, candidate_name: str | None, job: dict) -> str:
    """The {Name}_{DocType}_{Organization}_{hash}.docx format used briefly
    on 2026-08-05 (commit 51120a7), before Zahir asked for the trailing hex
    suffix to be removed entirely - kept only so
    _migrate_hash_suffixed_filename() can find and rename a folder's
    existing files under this now-superseded format."""
    doc_label = _sanitize_filename_part(DOC_LABELS.get(field, field))
    name = _sanitize_filename_part(candidate_name or "")
    org = _sanitize_filename_part(job.get("organization") or "")
    suffix = _job_hash(job.get("source"), job.get("job_id"), _WORKSPACE_FILENAME_SUFFIX_LEN)

    ext = ".docx"
    fixed_len = len(ext) + len("_" + suffix) + len("_") * 2
    budget = max(0, _MAX_WORKSPACE_FILENAME_LEN - fixed_len - len(doc_label))
    if len(name) + len(org) > budget:
        if len(name) <= budget // 2 or len(name) <= budget - len(org):
            name_budget = min(len(name), budget)
            org_budget = budget - name_budget
        elif len(org) <= budget // 2:
            org_budget = min(len(org), budget)
            name_budget = budget - org_budget
        else:
            name_budget = budget // 2
            org_budget = budget - name_budget
        name = name[:name_budget]
        org = org[:org_budget]

    parts = [p for p in [name, doc_label, org] if p]
    return "_".join(parts) + f"_{suffix}" + ext


def _workspace_filename(field: str, candidate_name: str | None, job: dict) -> str:
    """Filename for a workspace document: {Name}_{DocType}_{Organization}.docx,
    capped at _MAX_WORKSPACE_FILENAME_LEN total so it survives being pulled
    out of its folder into a destination with its own length limit (email,
    some cloud drives) - Title is dropped entirely (still preserved in the
    parent dossier folder name via _slug(), so nothing is lost, just not
    duplicated in the filename too). DocType is never truncated (it's what
    distinguishes Resume from CoverLetter); Name/Organization share
    whatever's left, and only the one that's actually long gives up length
    - if one is short, it keeps its full length and the long one absorbs
    the whole cut, rather than both being clipped by a flat amount.

    No disambiguating suffix (Zahir, 2026-08-05: didn't want a trailing hex
    string on the filename - see _legacy_hash_suffixed_workspace_filename()
    for the brief format this replaces). Within the app itself this can
    never actually collide: every job already gets its own physically
    separate dossier folder (_slug() below folds in a per-job hash), so two
    different jobs' documents are never written into the same directory in
    the first place - sync_workspace_documents()'s existing edit-detection/
    backup logic is what governs a filename recurring WITHIN one job's own
    folder (a normal regenerate), not a cross-job collision. The only way
    two different jobs could ever produce the identical
    Name_DocType_Organization.docx is if Zahir manually gathers files out of
    two different per-job folders into one flat destination himself
    (Downloads, an email) - accepted as a real but narrow and low-likelihood
    tradeoff (requires applying to 2+ different roles at the literal same
    company) rather than adding cross-folder collision-scanning complexity
    (an O(applications) scan on every document write) to guard against it."""
    doc_label = _sanitize_filename_part(DOC_LABELS.get(field, field))
    name = _sanitize_filename_part(candidate_name or "")
    org = _sanitize_filename_part(job.get("organization") or "")

    ext = ".docx"
    fixed_len = len(ext) + len("_") * 2  # 2 separators assuming name+org both present
    budget = max(0, _MAX_WORKSPACE_FILENAME_LEN - fixed_len - len(doc_label))
    if len(name) + len(org) > budget:
        if len(name) <= budget // 2 or len(name) <= budget - len(org):
            name_budget = min(len(name), budget)
            org_budget = budget - name_budget
        elif len(org) <= budget // 2:
            org_budget = min(len(org), budget)
            name_budget = budget - org_budget
        else:
            name_budget = budget // 2
            org_budget = budget - name_budget
        name = name[:name_budget]
        org = org[:org_budget]

    parts = [p for p in [name, doc_label, org] if p]
    return "_".join(parts) + ext


def _draft_unconfirmed_filename(real_filename: str) -> str:
    """The real filename with a "-DRAFT-UNCONFIRMED" marker inserted before
    the extension - deliberately NOT subject to _MAX_WORKSPACE_FILENAME_LEN's
    normal cap (that budget exists so a file survives being emailed/moved
    off Panga entirely; this filename is the opposite - a purposefully
    unmissable, internal-only marker that only ever needs to make sense
    inside this one dossier folder while unresolved, never meant to leave
    it)."""
    stem, ext = real_filename.rsplit(".", 1)
    return f"{stem}-DRAFT-UNCONFIRMED.{ext}"


def _migrate_legacy_filename(folder: Path, field: str, new_filename: str) -> None:
    """One-time rename for a workspace folder created before descriptive
    filenames existed: if the old generic name (resume.docx, etc.) is still
    on disk and the new descriptive name isn't, renames it in place so
    Zahir's real edits carry over under the new convention rather than
    being silently orphaned under the old name."""
    legacy_name = _LEGACY_WORKSPACE_FILENAMES.get(field)
    if not legacy_name:
        return
    legacy_path = folder / legacy_name
    new_path = folder / new_filename
    if legacy_path.exists() and not new_path.exists():
        legacy_path.rename(new_path)


def _migrate_descriptive_filename(folder: Path, field: str, new_filename: str, candidate_name: str | None, job: dict) -> None:
    """One-time rename for a workspace folder created under the 2026-07-31
    through 2026-08-05 {Name}_{DocType}_{Title}_{Organization}.docx format,
    before the 50-char ceiling existed: if that longer descriptive name is
    still on disk and the new short name isn't, renames it in place - same
    reasoning as _migrate_legacy_filename() one step further down the same
    chain, so Zahir's real edits under the older format aren't orphaned
    when the naming convention changes again."""
    old_name = _legacy_descriptive_workspace_filename(field, candidate_name, job)
    old_path = folder / old_name
    new_path = folder / new_filename
    if old_name != new_filename and old_path.exists() and not new_path.exists():
        old_path.rename(new_path)


def _migrate_hash_suffixed_filename(folder: Path, field: str, new_filename: str, candidate_name: str | None, job: dict) -> None:
    """One-time rename for a workspace folder created under the brief
    2026-08-05 {Name}_{DocType}_{Organization}_{hash}.docx format (commit
    51120a7), before Zahir asked for the trailing hex suffix to be dropped:
    if that hash-suffixed name is still on disk and the new suffix-free
    name isn't, renames it in place - same one-more-link-in-the-chain
    reasoning as _migrate_descriptive_filename(), so real edits made in the
    brief window that format existed aren't orphaned."""
    old_name = _legacy_hash_suffixed_workspace_filename(field, candidate_name, job)
    old_path = folder / old_name
    new_path = folder / new_filename
    if old_name != new_filename and old_path.exists() and not new_path.exists():
        old_path.rename(new_path)


def _format_pay(job: dict) -> str | None:
    if job.get("salary_text"):
        return job["salary_text"]
    if job.get("pay_min") or job.get("pay_max"):
        return f"${job.get('pay_min') or '?'}-${job.get('pay_max') or '?'}"
    return None


def write_dossier(source: str, job_id: str) -> Path | None:
    """Regenerates the dossier file for one job from current data. Returns
    None (writes nothing) if the job itself isn't in jobs.json - a dossier
    only makes sense once there's a real posting to attach it to."""
    job = next((j for j in load_jobs() if j.get("source") == source and j.get("job_id") == job_id), None)
    if job is None:
        return None
    application = get_application(source, job_id) or {}
    prep = get_interview_prep(source, job_id)

    lines = [
        f"# {job.get('title')} - {job.get('organization')}",
        "",
        f"- **Source:** {source}",
        f"- **Posting URL:** {job.get('posting_url') or 'n/a'}",
        f"- **Location:** {job.get('location') or 'not listed'}",
    ]
    pay = _format_pay(job)
    if pay:
        lines.append(f"- **Pay:** {pay}")
    if "fit_score" in job:
        lines.append(f"- **Compatibility score:** {job['fit_score']}/100")
        if job.get("fit_rationale"):
            lines.append(f"- **Why:** {job['fit_rationale']}")
    lines.append(f"- **Status:** {application.get('status') or 'not yet tracked'}")
    if application.get("strategy_tag"):
        lines.append(f"- **Strategy tag:** {application['strategy_tag']}")
    if application.get("skip_reason"):
        # The inline "Pass" dialog's "Something else" category stores
        # category+typed-detail "\n"-joined in one skip_reason string (see
        # ui/app.py's _pass_reason_dialog) - rendered as a nested bullet
        # here so it reads as two distinct lines instead of markdown's
        # lazy-continuation folding it into one run-on line. Every other
        # category (and the older single-string "Why not interested?" box)
        # has no embedded newline, so this is a no-op for those.
        reason_parts = application["skip_reason"].split("\n", 1)
        lines.append(f"- **Skip reason:** {reason_parts[0]}")
        if len(reason_parts) == 2:
            lines.append(f"  - {reason_parts[1]}")
    if application.get("created_at"):
        lines.append(f"- **Application started:** {application['created_at']}")
    if application.get("status_updated_at"):
        lines.append(f"- **Last status update:** {application['status_updated_at']}")
    lines.append("")

    if any(application.get(field) for field in DOC_LABELS):
        lines.append("## Documents")
        lines.append("")
        for field, label in DOC_LABELS.items():
            text = application.get(field)
            if text:
                lines.append(f"### {label}")
                lines.append("")
                lines.append(text)
                lines.append("")

    if prep and prep.get("rounds"):
        lines.append("## Interview prep")
        lines.append("")
        for round_ in prep["rounds"]:
            status_note = round_.get("outcome") or round_.get("status") or "in progress"
            lines.append(f"### {round_['round_label']} - {status_note}")
            logistics = " - ".join(v for v in [round_.get("date"), round_.get("format")] if v)
            if logistics:
                lines.append(f"*{logistics}*")
            lines.append("")
            for person in round_.get("interviewers") or []:
                title_part = f", {person['title']}" if person.get("title") else ""
                lines.append(f"**{person.get('name')}{title_part}**")
                if person.get("research_summary"):
                    lines.append(person["research_summary"])
                if person.get("persona"):
                    lines.append(f"_Likely focus:_ {person['persona']}")
                lines.append("")
            if round_.get("company_snapshot"):
                lines.append(f"**Company snapshot:** {round_['company_snapshot']}")
                lines.append("")
            if round_.get("likely_questions"):
                lines.append("**Likely questions:**")
                for q in round_["likely_questions"]:
                    asked_by = f" ({q['asked_by']})" if q.get("asked_by") else ""
                    lines.append(f"- {q.get('question')}{asked_by}")
                    if q.get("why"):
                        lines.append(f"  - why: {q['why']}")
                    if q.get("talking_point"):
                        lines.append(f"  - talking point: {q['talking_point']}")
                lines.append("")
            if round_.get("questions_to_ask"):
                lines.append("**Questions to ask them:**")
                for q in round_["questions_to_ask"]:
                    best_for = f" (best for {q['best_for']})" if q.get("best_for") else ""
                    lines.append(f"- {q.get('question')}{best_for}")
                lines.append("")
            if round_.get("outcome_notes"):
                lines.append(f"**Notes on how it went:** {round_['outcome_notes']}")
                lines.append("")

    path = dossier_path(source, job_id, job.get("organization"), job.get("title"))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def _extract_docx_text_from_document(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_docx_text(path: Path) -> str:
    # Deliberately not reused from profile/ingest.py's extract_text() -
    # same reasoning as linkedin/ingest.py's separate _extract_pdf_text():
    # the logic is small and the packages serve different concerns, so a
    # cross-package dependency isn't worth it for ~10 lines.
    return _extract_docx_text_from_document(Document(str(path)))


def _render_doc_key_to_bytes(doc_key: str, text: str, candidate_name: str | None, job: dict) -> bytes:
    """The same doc_key -> renderer/params dispatch sync_workspace_documents()
    uses to write a fresh .docx, factored out so edit-detection can render
    the SAME stored text through the SAME pipeline before comparing it
    against what's on disk - see _baseline_docx_text()'s docstring for why
    this matters."""
    if doc_key == "cover_letter":
        return cover_letter_to_docx_bytes(
            text, company_name=job.get("organization"),
            company_address=job.get("organization_address"), author=candidate_name,
        )
    if doc_key == "resume" and job.get("source") == "USAJOBS":
        word_count = len(text.split())
        body_size_pt = 9.0 if word_count > 1100 else 10.0
        return text_to_docx_bytes(text, author=candidate_name, body_size_pt=body_size_pt)
    return text_to_docx_bytes(text, author=candidate_name)


def _baseline_docx_text(doc_key: str, text: str, candidate_name: str | None, job: dict) -> str:
    """What _extract_docx_text() would return for `text` if it had just been
    (re)rendered and never touched since - the correct thing to diff a
    workspace file against, NOT the raw stored text itself.

    Real bug fixed 2026-08-04 (Zahir hit this live: a cover letter he never
    touched was reported as "4 lines added, 6 removed"): the renderers
    inject content that never existed in the raw drafted text at all
    (cover_letter_to_docx_bytes always adds today's date, the company name,
    and an address line before the body) and also transform it on the way
    in (blank lines dropped, "- " bullet markers stripped in favor of the
    "List Bullet" paragraph style). Comparing raw stored text against
    text extracted back out of the rendered file meant EVERY document with
    a bullet, or EVERY cover letter at all, showed a spurious diff on every
    check - not just here in check_for_edits(), but also in
    sync_workspace_documents()'s own edit-detection, which was silently
    creating an unnecessary ".edited-<timestamp>.docx" backup on every
    single regenerate, edited or not."""
    return _extract_docx_text_from_document(
        Document(io.BytesIO(_render_doc_key_to_bytes(doc_key, text, candidate_name, job)))
    )


def sync_workspace_documents(
    source: str,
    job_id: str,
    doc_keys: list[str],
    drafted: dict,
    profile: dict,
    job: dict,
) -> None:
    """Writes the real, editable .docx files for exactly the doc_keys just
    (re)generated by tailoring.drafting.generate_documents() - called ONLY
    right after that succeeds (see ui/app.py), never from write_dossier()'s
    generic refresh path. This is deliberate: Zahir is meant to open these
    files directly in Word and edit them, so anything that rewrote them on
    an unrelated save (a status change, a strategy tag) would silently
    destroy his in-progress edits.

    If a workspace file already exists AND its current on-disk text differs
    from what's already stored in applications.json for that field (i.e.
    Zahir appears to have edited his working copy since the last draft),
    the existing file is renamed to a timestamped backup
    (resume.edited-2026-07-31T121500.docx) before the fresh draft is
    written - his edits are preserved on disk, never silently overwritten,
    even though the app can no longer treat that specific text as "the
    current working copy" once a fresh regenerate has happened.

    2026-08-09, the real safety guarantee behind the unconfirmed-"?"-claims
    feature: if a field's new text still has an unresolved "?" in it
    (tailoring.unconfirmed_claims.has_unconfirmed_marker), this writes to a
    distinctly-named *-DRAFT-UNCONFIRMED.docx file instead of the real
    expected filename, and leaves any existing real file (from a
    previously resolved draft) untouched rather than overwriting it with
    unconfirmed content. This is what makes it safe to write a hedged "?"
    guess into resume_text at all - real gap General/Zahir caught before
    shipping: blocking only in-app actions (the "applied" status
    transition, the Apply Assist render) still leaves the actual .docx
    sitting on disk, fully double-clickable/attachable from File Explorer
    regardless of any in-app gate - so the gate has to live here, at the
    one place that ever writes the real filename, not just upstream of it.
    Once a field's text is clean again, any stale *-DRAFT-UNCONFIRMED.docx
    from an earlier unresolved round is deleted - it served its purpose
    once the real file exists for real.
    apply_answers is skipped here (see DOC_KEY_TO_FIELD) - it isn't rendered
    as an editable docx, the Results tab shows it as copy-paste fields
    directly.
    """
    doc_keys_with_files = [k for k in doc_keys if k in DOC_KEY_TO_FIELD]
    if not doc_keys_with_files:
        return
    application = get_application(source, job_id) or {}
    folder = dossier_dir(source, job_id, job.get("organization"), job.get("title"))
    folder.mkdir(parents=True, exist_ok=True)
    candidate_name = profile.get("name")

    for doc_key in doc_keys_with_files:
        field = f"{doc_key}_text"
        new_text = drafted.get(doc_key)
        if doc_key == "resume" and isinstance(new_text, dict):
            new_text = new_text.get("text")
        if not new_text:
            continue

        filename = _workspace_filename(field, candidate_name, job)
        _migrate_legacy_filename(folder, field, filename)
        _migrate_descriptive_filename(folder, field, filename, candidate_name, job)
        _migrate_hash_suffixed_filename(folder, field, filename, candidate_name, job)
        file_path = folder / filename
        draft_path = folder / _draft_unconfirmed_filename(filename)

        if has_unconfirmed_marker(new_text):
            # Real gap General/Zahir caught before this shipped: blocking
            # only the "applied" status transition and the Apply Assist
            # render still left the REAL synced .docx sitting on disk with
            # an unresolved "?" in it, fully double-clickable/attachable
            # from File Explorer regardless of either in-app gate. This is
            # the actual guarantee: the real expected filename is never
            # created/overwritten from text that still has a "?" in it -
            # written to a distinctly-named draft file instead, and the
            # last known-good real file (if one exists from a previously
            # resolved draft) is left untouched rather than overwritten
            # with unconfirmed content.
            doc_bytes = _render_doc_key_to_bytes(doc_key, new_text, candidate_name, job)
            draft_path.write_bytes(doc_bytes)
            continue

        if file_path.exists():
            previous_stored_text = application.get(field)
            try:
                current_on_disk_text = _extract_docx_text(file_path)
            except Exception:
                current_on_disk_text = None
            # Compare against a freshly-rendered baseline of the PREVIOUS
            # stored text, not the raw text itself - see _baseline_docx_text()
            # for why a raw-text comparison here was a real bug (every
            # regenerate looked like an edit, even when the file was never
            # touched, spamming spurious .edited-*.docx backups).
            baseline_text = _baseline_docx_text(doc_key, previous_stored_text or "", candidate_name, job)
            if current_on_disk_text is not None and current_on_disk_text.strip() != baseline_text.strip():
                stamp = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H%M%S")
                backup_name = f"{file_path.stem}.edited-{stamp}{file_path.suffix}"
                file_path.rename(folder / backup_name)

        doc_bytes = _render_doc_key_to_bytes(doc_key, new_text, candidate_name, job)
        file_path.write_bytes(doc_bytes)
        # This text is clean (no gate above triggered) - a stale draft
        # file from a previously unresolved round, if any, is no longer
        # needed once the real file has been written for real.
        if draft_path.exists():
            draft_path.unlink()


def check_for_edits(source: str, job_id: str) -> dict:
    """Reads back each workspace document that exists on disk and diffs it
    against the text currently stored in applications.json for that field -
    this is how Panga detects that Zahir edited his working copy, without
    needing him to paste anything back manually. Returns
    {doc_key: {"changed": bool, "diff": [...]}} for every doc_key with
    stored text, one entry per drafted document regardless of whether a
    workspace file exists yet. A document drafted before this workspace
    feature existed (or not regenerated since) has stored text but no
    on-disk file yet - that case returns {"changed": False, "diff": [],
    "no_workspace_file": True} rather than being silently omitted, so the
    Results tab can tell "nothing changed" apart from "there's nothing to
    compare against yet, regenerate this document to create it." diff is a
    compact unified-diff line list (empty if unchanged or no file)."""
    application = get_application(source, job_id) or {}
    job = next((j for j in load_jobs() if j.get("source") == source and j.get("job_id") == job_id), None)
    if job is None:
        return {}
    folder = dossier_dir(source, job_id, job.get("organization"), job.get("title"))
    candidate_name = load_profile().get("name")

    results = {}
    for doc_key, field in DOC_KEY_TO_FIELD.items():
        stored_text = application.get(field)
        if not stored_text:
            continue
        filename = _workspace_filename(field, candidate_name, job)
        _migrate_legacy_filename(folder, field, filename)
        _migrate_descriptive_filename(folder, field, filename, candidate_name, job)
        _migrate_hash_suffixed_filename(folder, field, filename, candidate_name, job)
        file_path = folder / filename
        if not file_path.exists():
            # A document drafted before this workspace feature existed (or
            # never re-generated since) has stored text but no editable
            # file yet - distinct from "checked, no changes found", which
            # is what an empty diff on an EXISTING file means below.
            results[doc_key] = {"changed": False, "diff": [], "no_workspace_file": True}
            continue
        try:
            current_text = _extract_docx_text(file_path)
        except Exception as exc:
            results[doc_key] = {"changed": False, "diff": [], "error": f"Couldn't read {filename}: {exc}"}
            continue
        # Diff against a freshly-rendered baseline of stored_text, not
        # stored_text itself - see _baseline_docx_text() for why comparing
        # raw drafted text against text extracted back out of a rendered
        # .docx was a real bug (every cover letter, and every document with
        # a bullet, showed a spurious diff on every check).
        baseline_text = _baseline_docx_text(doc_key, stored_text, candidate_name, job)
        if current_text.strip() == baseline_text.strip():
            results[doc_key] = {"changed": False, "diff": []}
        else:
            diff = list(difflib.unified_diff(
                baseline_text.splitlines(), current_text.splitlines(),
                fromfile="original draft", tofile="your edited version", lineterm="",
            ))
            results[doc_key] = {"changed": True, "diff": diff}
    return results
