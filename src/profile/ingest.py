"""Build step 1: parses source Word documents (per config/document_manifest.yaml)
into raw text, saved locally under data/profile/raw/. This is a text-extraction
pass only — turning the extracted text into structured fields (roles, dates,
skills) happens via the gap-probing interview and later reasoning, not here.
Raw text is encrypted at rest (PRD §7) via security.crypto_store.
"""

from pathlib import Path

import yaml
from docx import Document
from pypdf import PdfReader

import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from security.crypto_store import read_json, read_text, write_text, write_json  # noqa: E402

PROJECT_ROOT = Path(__file__).resolve().parents[2]
MANIFEST_PATH = PROJECT_ROOT / "config" / "document_manifest.yaml"
OUTPUT_DIR = PROJECT_ROOT / "data" / "profile" / "raw"
MANIFEST_RESULT_PATH = OUTPUT_DIR / "manifest_result.json"


def extract_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def slugify(filename: str) -> str:
    stem = Path(filename).stem
    return "".join(c if c.isalnum() else "_" for c in stem).strip("_").lower()


def ingest_all() -> list[dict]:
    manifest = yaml.safe_load(MANIFEST_PATH.read_text(encoding="utf-8"))
    source_folder = Path(manifest["source_folder"])

    results = []
    for entry in manifest["documents"]:
        source_path = source_folder / entry["file"]
        text = extract_text(source_path)
        slug = slugify(entry["file"])
        out_path = OUTPUT_DIR / f"{slug}.txt"
        write_text(out_path, text)

        results.append({
            "source_file": entry["file"],
            "category": entry["category"],
            "target_title": entry.get("target_title"),
            "note": entry.get("note"),
            "word_count": len(text.split()),
            "extracted_to": str(out_path.relative_to(PROJECT_ROOT)),
        })

    write_json(MANIFEST_RESULT_PATH, results)
    return results


def _extract_pdf_text(file) -> str:
    reader = PdfReader(file)
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in parts if p.strip())


def load_manifest_result() -> list[dict]:
    return read_json(MANIFEST_RESULT_PATH, default=[])


def resume_text() -> str:
    """Concatenated raw text of every ingested document tagged
    category=="resume" - shared by profile.interview's gap detection and
    tailoring.drafting.generate_target_roles(), so both reason over the
    same source text instead of each re-reading manifest_result.json."""
    chunks = []
    for entry in load_manifest_result():
        if entry["category"] == "resume":
            chunks.append(read_text(PROJECT_ROOT / entry["extracted_to"]))
    return "\n".join(chunks)


def remove_document(source_file: str) -> None:
    """Deletes the extracted raw-text file and drops the manifest entry for
    it. Filter-and-resave, not blank-and-resave - matches every other store
    in this project (write_json always takes the whole list)."""
    entries = load_manifest_result()
    entry = next((e for e in entries if e["source_file"] == source_file), None)
    if entry:
        (PROJECT_ROOT / entry["extracted_to"]).unlink(missing_ok=True)
    write_json(MANIFEST_RESULT_PATH, [e for e in entries if e["source_file"] != source_file])


def ingest_uploaded_document(
    file, filename: str, category: str, target_title: str | None = None, note: str | None = None,
) -> dict:
    """Streamlit-uploader counterpart to ingest_all() - takes an uploaded
    file directly instead of reading source_folder from
    config/document_manifest.yaml, so Zahir can add/replace a document
    himself without a Claude session editing that manifest by hand. Same
    output shape (manifest_result.json + a raw-text file under
    data/profile/raw/), so interview.py's category=="resume" read keeps
    working unchanged. Re-uploading a file with the same name replaces its
    existing entry rather than duplicating it."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        text = extract_text(file)
    elif ext == ".pdf":
        text = _extract_pdf_text(file)
    else:
        raise ValueError(f"Unsupported file type \"{ext}\" - upload a .docx or .pdf.")

    slug = slugify(filename)
    out_path = OUTPUT_DIR / f"{slug}.txt"
    write_text(out_path, text)

    entry = {
        "source_file": filename,
        "category": category,
        "target_title": target_title,
        "note": note,
        "word_count": len(text.split()),
        "extracted_to": str(out_path.relative_to(PROJECT_ROOT)),
    }
    entries = [e for e in load_manifest_result() if e["source_file"] != filename]
    entries.append(entry)
    write_json(MANIFEST_RESULT_PATH, entries)
    return entry


if __name__ == "__main__":
    for r in ingest_all():
        print(f"{r['category']:9} {r['word_count']:5} words  <- {r['source_file']}")
